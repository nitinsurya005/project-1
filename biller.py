from docx import *
from PIL import ImageTk, Image
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from tkinter import *
import os
from docx2pdf import *
import sys

class biller:
    def __init__(self,list,o,clr=1,mode=1,save_l='F:\\Bills',m=1):
        self.save_l=save_l
        self.clr=clr
        self.invno=list[0]
        self.over=False
        import pythoncom
        pythoncom.CoInitialize()
        self.o=os.getcwd() if m else o
        self.o=self.o.rstrip('\\Files')
        # self.o='G:\\PROJECTS\\Python\\PROJECTS\\Unfinished\\Invoice Maker'
        self.o=o+'\\'
        self.over=False
        list.insert(0,'Copy')
        if mode==1:
            file='Sales'
            cl=[0, 21, 27, 49, 55, 77, 83, 99, 105, 111, 133, 139, 161, 167, 183, 189, 211, 212, 217, 218, 221, 222, 224, 238, 253, 283, 285, 297, 299, 311, 313, 327]
            nof='nos'
            ol=[99,183]
        else:
            file='Purchase'
            cl=[0, 14, 19, 34, 39, 71, 91, 92, 94, 95, 97, 98, 100, 110, 121, 133]
            nof='nop'
            ol=[71,133]
        fname=['Original',"Record's Copy","Transporter's Copy"]
        cname=['(Original for Recipient)','(Duplicate for Records)','(Duplicate for Transporter)']
        for j in range(3):
            self.over=False
            i=0
            ii=0
            list[0]=cname[j]
            wordDoc=Document(self.o+f'Files\\{file}.docx')
            for table in wordDoc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if i in cl:
                            cell.text=str(list[ii])
                            p=cell.paragraphs[0]
                            if i not in ol:
                                p=cell.paragraphs[0]
                                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                                cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
                            run=p.runs
                            font=run[0].font
                            font.size=Pt(13)
                            ii+=1
                        i+=1
            sname=f'{self.save_l}\\{file}\\{fname[j]}\\Bill {self.invno} {fname[j]}'
            wordDoc.save(f'{sname}.docx')
            if mode!=3:
                convert(f'{sname}.docx',f'{sname}.pdf')
                os.remove(f'{sname}.docx')
            
        if mode in (1,2) and m==1:
            a=open(self.o.rstrip('\\Files')+f'\\Files\\{nof}.txt','r')
            b=a.read()
            a.close()
            if str(self.invno)==str(b):
                x=open(self.o.rstrip('\\Files')+f'\\Files\\{nof}.txt','w')
                x.write(str(int(b)+1))
                x.close()
                self.invno=str(int(b)+1)
            else:
                self.invno=b
            self.clr()
            self.over=True
            root2=Tk()
            root2.title('Successful!')
            root2['bg']='#ffffff'
            c=Label(root2,text='Bill Created Successfully',bg='#ffffff',font=('comic sans ms',15,'bold')).pack()
            cb=Button(root2,text='Close',command=root2.destroy,font=('comic sans ms',15,'bold'),fg='#000000',bg='#ffffff').pack()
            root2.mainloop()